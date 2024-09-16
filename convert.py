from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import markdown
from bs4 import BeautifulSoup

def add_heading(doc, text, level):
    """Add a heading to the docx document."""
    doc.add_heading(text, level=level)

def add_paragraph(doc, text, bold=False, italic=False, indent=False):
    """Add a formatted paragraph to the docx document."""
    para = doc.add_paragraph()
    if indent:
        para.paragraph_format.left_indent = Pt(24)  # Adjust indent size for blockquote
    run = para.add_run(text)
    if bold:
        run.bold = True
    if italic:
        run.italic = True

def add_list(doc, items, ordered=False):
    """Add a list (ordered or unordered) to the docx document."""
    for index, item in enumerate(items):
        if ordered:
            doc.add_paragraph(f"{index + 1}. {item}", style='List Number')
        else:
            doc.add_paragraph(f"• {item}", style='List Bullet')

def convert_md_to_docx(md_file_path, docx_file_path):
    # Open the Markdown file and convert it to HTML
    with open(md_file_path, 'r', encoding='utf-8', errors='ignore') as f:
        md_content = f.read()

    # Convert markdown to HTML
    html_content = markdown.markdown(md_content)

    # Parse HTML to extract elements
    soup = BeautifulSoup(html_content, 'html.parser')

    # Create a new docx Document
    doc = Document()

    # Parse the HTML and add corresponding docx elements
    for element in soup:
        if element.name == 'blockquote':
            # Blockquote: Indent the text
            for child in element.children:
                if child.name == 'h2':
                    add_heading(doc, child.get_text(), level=2)
                elif child.name == 'p':
                    add_paragraph(doc, child.get_text(), indent=True)
        elif element.name == 'h2':
            add_heading(doc, element.get_text(), level=2)
        elif element.name == 'h3':
            add_heading(doc, element.get_text(), level=3)
        elif element.name == 'p':
            # Handle bold and italic within paragraphs
            bold_elements = element.find_all('strong')
            italic_elements = element.find_all('em')
            
            if bold_elements or italic_elements:
                for sub_element in element.children:
                    if sub_element.name == 'strong':
                        add_paragraph(doc, sub_element.get_text(), bold=True)
                    elif sub_element.name == 'em':
                        add_paragraph(doc, sub_element.get_text(), italic=True)
                    else:
                        add_paragraph(doc, sub_element.get_text())
            else:
                add_paragraph(doc, element.get_text())
        elif element.name == 'ul':
            # Unordered list
            list_items = [li.get_text() for li in element.find_all('li')]
            add_list(doc, list_items, ordered=False)
        elif element.name == 'ol':
            # Ordered list
            list_items = [li.get_text() for li in element.find_all('li')]
            add_list(doc, list_items, ordered=True)
        elif element.name == 'a':
            # Hyperlink
            link_text = element.get_text()
            link_url = element['href']
            add_paragraph(doc, f'{link_text} ({link_url})')
        elif element.name == 'img':
            # Images
            img_src = element['src']
            alt_text = element.get('alt', 'Image')
            add_paragraph(doc, f'[Image: {alt_text}, URL: {img_src}]')

    # Save the document
    doc.save(docx_file_path)
    print(f"Converted '{md_file_path}' to '{docx_file_path}' successfully.")

# Example usage
if __name__=='__main__':
    md_file = "example.md"   # Path to the Markdown file
    docx_file = "output.docx"  # Path to the output docx file
    convert_md_to_docx(md_file, docx_file)
